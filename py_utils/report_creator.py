'''
Created on 3 de jun. de 2016

@author: daniel.cano
Descripcion:
busca el CSV, genera la imagen de la grafica de posicion capturada
por el util de verificacion.
Luego genera un word con todas las imagenes
'''
import os
from docx import Document

filename = ''
# path = 'D:/Workspace/Verificaciones/16_06_01_GreatPlains/'
# path = 'D:/Workspace/Verificaciones/GP_Speed_range/'
path = 'D:/Workspace/Verificaciones/16_07_04_ADT/'

'''search for .csv and generate graphs'''
for root, dirs, files in os.walk(path):
    for name in files:
        if name.endswith(("D.csv")):
            args = '-p ' + path + ' -f '+ name
            os.system("speed_calculator_from_excel.py %s " % ''.join(args))

'''create document and insert the graphs'''
document = Document()
paragraph = document.add_paragraph('Data results graphs:')


for root, dirs, files in os.walk(path+'/images/'):
    for name in files:
        if name.endswith((".png")):
            document.add_picture(path+'/images/'+name)

document.save(path + 'test.docx')

print '*** END PROGRAM ***'
